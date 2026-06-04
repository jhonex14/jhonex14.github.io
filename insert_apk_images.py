import os
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def insert_apk_images():
    doc_path = r"C:\Users\jhone\Desktop\CONSULTIME-THESIS FINAL NA EDIT.docx"
    if not os.path.exists(doc_path):
        print(f"Error: '{doc_path}' not found.")
        return

    doc = docx.Document(doc_path)
    print(f"Loaded document with {len(doc.paragraphs)} paragraphs.")

    def add_image_after(para_keyword, image_path, width_in, caption_text):
        if not os.path.exists(image_path):
            print(f"Warning: image '{image_path}' not found. Skipping.")
            return False
            
        target_p = None
        for i, p in enumerate(doc.paragraphs):
            if para_keyword in p.text:
                target_p = p
                print(f"Found target paragraph index {i}: '{p.text[:60]}...'")
                break
                
        if target_p:
            # Add paragraph for image
            new_p = doc.add_paragraph()
            target_p._element.addnext(new_p._element)
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = new_p.add_run()
            run.add_picture(image_path, width=Inches(width_in))
            
            # Add caption paragraph
            cap_p = doc.add_paragraph()
            new_p._element.addnext(cap_p._element)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cap_run = cap_p.add_run(caption_text)
            cap_run.italic = True
            cap_run.font.size = Pt(10)
            cap_run.font.name = 'Times New Roman'
            
            cap_p.paragraph_format.space_before = Pt(4)
            cap_p.paragraph_format.space_after = Pt(12)
            
            print(f"SUCCESS: Inserted '{image_path}' after keyword")
            return True
        else:
            print(f"Error: Could not find paragraph with keyword '{para_keyword}'")
            return False

    # Insert APK Download Button Screenshot
    add_image_after(
        "A premium, outline-border Android APK download button has been added directly to both the login and registration portals.",
        "apk_download_button.png",
        4.5,
        "Figure 3.3. ConsulTime Client UI integrated Direct Android App (.APK) download trigger button."
    )

    # Insert APK MediaFire Page Screenshot
    add_image_after(
        "allowing the application to be deployed both as an online website and a downloadable Android .APK file.",
        "apk_mediafire_download.png",
        5.5,
        "Figure 2.4. MediaFire host server interface distributing the compiled 3.69MB native application installer package (consultime.apk)."
    )

    doc.save(doc_path)
    print("SUCCESS: Finished inserting all APK-related screenshots into the thesis Word document!")

if __name__ == "__main__":
    insert_apk_images()
