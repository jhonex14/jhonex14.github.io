import os
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def insert_images():
    doc_path = r"C:\Users\jhone\Desktop\CONSULTIME-THESIS FINAL NA EDIT.docx"
    if not os.path.exists(doc_path):
        print(f"Error: '{doc_path}' not found.")
        return

    doc = docx.Document(doc_path)
    print(f"Successfully loaded document with {len(doc.paragraphs)} paragraphs.")

    # Helper function to insert an image and caption after a target paragraph
    def add_image_after(para_keyword, image_path, width_in, caption_text):
        if not os.path.exists(image_path):
            print(f"Warning: image '{image_path}' not found. Skipping.")
            return False
            
        target_p = None
        for i, p in enumerate(doc.paragraphs):
            if para_keyword in p.text:
                target_p = p
                print(f"Found target paragraph index {i}: '{p.text[:60]}...'")
                break
                
        if target_p:
            # Add new paragraph for image
            new_p = doc.add_paragraph()
            target_p._element.addnext(new_p._element)
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = new_p.add_run()
            run.add_picture(image_path, width=Inches(width_in))
            
            # Add caption paragraph directly after image paragraph
            cap_p = doc.add_paragraph()
            new_p._element.addnext(cap_p._element)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cap_run = cap_p.add_run(caption_text)
            cap_run.italic = True
            cap_run.font.size = Pt(10)
            cap_run.font.name = 'Times New Roman'
            
            # Spacing between paragraphs
            cap_p.paragraph_format.space_before = Pt(4)
            cap_p.paragraph_format.space_after = Pt(12)
            
            print(f"SUCCESS: Inserted '{image_path}' after '{para_keyword}'")
            return True
        else:
            print(f"Error: Could not find paragraph containing '{para_keyword}'")
            return False

    # 1. Database schema screenshots
    add_image_after(
        "Administrative Governance: The inclusion of the boolean switch is_approved operates",
        "db_profiles_screenshot.png",
        6.0,
        "Figure 2.1. Supabase Profiles Database Table Schema representing user credentials and vetting switches."
    )
    
    add_image_after(
        "Operational Adaptability: The inclusion of an optional specific_date field provides",
        "db_availability_screenshot.png",
        6.0,
        "Figure 2.2. Supabase Faculty Availabilities Database Table Schema mapping recurring and specific date time blocks."
    )
    
    add_image_after(
        "Real-Time State Transitions: System states are managed through the strict status text",
        "db_bookings_screenshot.png",
        6.0,
        "Figure 2.3. Supabase Consultation Bookings Junction Database Table Schema tracking appointment states."
    )

    # 2. Admin dashboard & vetting screenshots
    add_image_after(
        "Figure 3. Multi-Role System Use Case, State Transitions",
        "admin_overview_screenshot.png",
        6.0,
        "Figure 3.1. Admin Portal dashboard overview interface monitoring system telemetry and quick action requests."
    )
    
    add_image_after(
        "Figure 3.1. Admin Portal dashboard overview",
        "admin_vetting_screenshot.png",
        6.0,
        "Figure 3.2. Admin Portal Faculty Approvals interface showing pending credentials vetting requests."
    )

    # 3. Cross-device responsiveness mockups
    add_image_after(
        "Vivo Y17s, 6.56\" (720x1612 px)",
        "consultime_mobile_mockup.png",
        5.5,
        "Figure 4.1. ConsulTime interface responsiveness on mobile device form factors (Vivo Y17s & Samsung Galaxy S9)."
    )
    
    add_image_after(
        "iPad Air Pro 11\" (2360x1640px)",
        "consultime_tablet_mockup.png",
        5.5,
        "Figure 4.2. ConsulTime interface responsiveness on tablet device form factors (iPad Air Pro)."
    )
    
    add_image_after(
        "Desktop Computer 24\" (1920x1080px)",
        "consultime_desktop_mockup.png",
        6.0,
        "Figure 4.3. ConsulTime interface responsiveness and grid density layout on laptop and desktop workstation screen resolutions."
    )

    doc.save(doc_path)
    print("SUCCESS: Finished inserting all screenshots into the thesis Word document!")

if __name__ == "__main__":
    insert_images()
